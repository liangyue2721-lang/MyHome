import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_full_content_resume():
    doc = Document()

    # --- 1. 全局页面设置 ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 颜色定义
    THEME_COLOR_HEX = "2F3A45" # 深铁灰
    BORDER_COLOR_HEX = "808080" # 灰色横线
    TEXT_BLACK = RGBColor(0, 0, 0)

    # --- 辅助样式函数 ---
    def set_font(run, size=10.5, bold=False, color=TEXT_BLACK):
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = bold
        run.font.color.rgb = color

    def set_cell_shading(cell, color_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def set_cell_bottom_border(cell, color_hex="000000", sz="8"):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), sz)
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), color_hex)
        tcBorders.append(bottom)

    def add_section_header(text):
        """左侧深色色块 + 右侧贯穿横线"""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(1.2)
        table.columns[1].width = Inches(6.0)

        row = table.rows[0]
        # 左侧色块
        c1 = row.cells[0]
        set_cell_shading(c1, THEME_COLOR_HEX)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = p1.add_run(text)
        set_font(run1, size=11, bold=True, color=RGBColor(255, 255, 255))
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)

        # 右侧横线
        c2 = row.cells[1]
        set_cell_bottom_border(c2, color_hex=BORDER_COLOR_HEX, sz="8")

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_3col_row(col1, col2, col3, is_bold=True):
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        table.columns[0].width = Inches(1.5) # 时间
        table.columns[1].width = Inches(4.0) # 公司/项目
        table.columns[2].width = Inches(1.7) # 职位

        row = table.rows[0]
        p1 = row.cells[0].paragraphs[0]
        set_font(p1.add_run(col1), size=10.5, bold=False)

        p2 = row.cells[1].paragraphs[0]
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_font(p2.add_run(col2), size=11, bold=is_bold)

        p3 = row.cells[2].paragraphs[0]
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        set_font(p3.add_run(col3), size=10.5, bold=is_bold)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        set_font(run, size=10.5)

    # ================= 简历正文 =================

    # 1. 姓名
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run("梁 朋")
    set_font(title_run, size=22, bold=True, color=RGBColor(47, 58, 69))
    title_p.paragraph_format.space_after = Pt(10)

    # 2. 基本信息
    add_section_header("基本信息")
    info_table = doc.add_table(rows=3, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Inches(3.5)
    info_table.columns[1].width = Inches(3.5)

    r1 = info_table.rows[0]
    set_font(r1.cells[0].paragraphs[0].add_run("性    别：男 | 1994.10"))
    set_font(r1.cells[1].paragraphs[0].add_run("学    历：河北机电职业技术学院 | 专科"))

    r2 = info_table.rows[1]
    set_font(r2.cells[0].paragraphs[0].add_run("电    话：15511644036"))
    set_font(r2.cells[1].paragraphs[0].add_run("邮    箱：15511644036@163.com"))

    r3 = info_table.rows[2]
    set_font(r3.cells[0].paragraphs[0].add_run("微    信：18330900451"))
    set_font(r3.cells[1].paragraphs[0].add_run("现    居：北京"))
    doc.add_paragraph()

    # 3. 专业技能
    add_section_header("专业技能")
    skills = [
        "Java 核心：扎实基础，熟悉多线程、JVM 内存模型及 GC 机制，具备线上问题排查能力。",
        "微服务架构：熟练使用 Spring Boot / Spring Cloud，具备丰富的服务治理与分布式架构经验。",
        "中间件：深入应用 Kafka/RabbitMQ 异步解耦；精通 Redis 缓存设计与热点优化。",
        "数据库：精通 MySQL 索引设计、慢 SQL 分析与性能调优。",
        "运维部署：熟悉 Linux/Docker/Nginx，具备生产环境部署与信创适配经验。"
    ]
    for s in skills:
        add_bullet(s)
    doc.add_paragraph()

    # 4. 工作经历 (已补充详细内容)
    add_section_header("工作经历")

    # --- Job 1 ---
    add_3col_row("2024.10 - 至今", "北京汉克时代科技有限公司", "Java 开发工程师")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.add_run("驻场：中信百信银行"), size=10, bold=True, color=RGBColor(80, 80, 80))
    # 补充职责
    job1_details = [
        "负责百信银行核心业务系统的持续迭代与维护，确保系统在生产环境的高可用性。",
        "参与系统性能监控与故障排查，针对慢SQL及热点数据问题进行专项优化。",
        "配合行方进行技术架构升级，推进系统稳定性建设。"
    ]
    for d in job1_details:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Job 2 ---
    add_3col_row("2021.08 - 2024.10", "软通动力信息技术（集团）有限公司", "Java 开发工程师")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.add_run("驻场：中信百信银行"), size=10, bold=True, color=RGBColor(80, 80, 80))
    # 补充职责 (关联项目1和2)
    job2_details = [
        "主导【统一文件传输平台(FTS)】的后端开发，支撑全行25+系统的大规模文件交互。",
        "负责【消息中心系统】的分布式架构设计，利用Kafka实现多渠道消息的异步削峰填谷。",
        "参与信创环境适配改造，完成国产化数据库及中间件的迁移与调优工作。"
    ]
    for d in job2_details:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Job 3 ---
    add_3col_row("2019.10 - 2021.08", "北京建设信源资讯有限公司", "Java 开发工程师")
    # 补充职责 (关联项目3)
    job3_details = [
        "负责【中信银行股东关联交易管理系统】的核心模块开发，对接董办及监管报送需求。",
        "实现疑似关联方识别算法（DFS股权穿透），解决复杂的股权关系计算难题。",
        "负责与行内多系统（HR、信贷、CRM）的数据集成与接口开发，确保数据一致性。"
    ]
    for d in job3_details:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Job 4 ---
    add_3col_row("2018.07 - 2019.10", "北京英源科技有限公司", "Java 开发工程师")
    job4_details = [
        "参与公司金融类系统功能模块的代码编写与单元测试。",
        "负责系统日常缺陷修复（Bug Fix）及文档编写，协助上线部署。"
    ]
    for d in job4_details:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 5. 核心项目经验 (项目顺序保持不变：FTS -> 消息 -> 关联交易)
    add_section_header("项目经历")

    # --- Project 1: FTS ---
    add_3col_row("项目一", "百信银行统一文件传输平台 (FTS)", "核心负责人")
    details_1 = [
        "核心架构开发：主导文件上传下载、路由分发、断点续传及分片传输等关键能力开发。",
        "高并发与性能优化：设计流量控制机制，引入 Kafka 异步解耦流水日志，降低数据库压力，提升吞吐量。",
        "安全与稳定性：实现基于 MD5 的文件完整性校验与分片加密传输；参与集群高可用部署及信创环境适配。"
    ]
    for d in details_1:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Project 2: 消息中心 ---
    add_3col_row("项目二", "消息中心系统", "后端开发")
    details_2 = [
        "异步削峰：基于 Kafka 实现消息的异步投递，彻底剥离消息服务与主业务流程的同步耦合。",
        "性能提升：利用 Redis 缓存消息模板与发送状态，大幅提升消息处理效率；设计失败重试机制，保障消息必达。"
    ]
    for d in details_2:
        add_bullet(d)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Project 3: 股东关联 ---
    add_3col_row("项目三", "股东关联交易管理系统（二期）", "核心开发")
    # 技术栈
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    set_font(p.add_run("技术栈：Spring Boot, MyBatis, MySQL, Redis, DFS 算法, 外部工商数据接口"), size=10, bold=True)

    details_3 = [
        "疑似关联方识别算法：负责核心识别模块开发，对接外部工商数据，基于深度优先搜索(DFS)算法实现股权穿透计算，自动识别“控制”、“共同控制”及“重大影响”下的疑似关联方。",
        "多系统深度集成：对接行内 HR 系统同步高管及亲属信息；对接新一代授信系统(520)与 CRM 系统，提供实时关联方查询接口，实现授信业务中的关联交易自动拦截与会签触发。",
        "数据报备工作流：设计从任务分发（总行/分行/子公司）、多级复核到汇总签批的全流程引擎，支持退回与动态流转，提升全行报送效率。",
        "报表与风控：实现关联交易金额的上限管控与实时预警；基于 POI 开发复杂监管报表（授信/非授信明细）的自动化生成。"
    ]
    for d in details_3:
        add_bullet(d)

    # 6. 自我评价
    add_section_header("自我评价")
    p_eval = doc.add_paragraph()
    text_eval = "工作积极认真，拥有 8 年 Java 后端开发经验，深耕互联网金融领域。擅长处理高并发、复杂业务逻辑及多系统集成问题。具备良好的团队协作能力与抗压能力。"
    set_font(p_eval.add_run(text_eval), size=10.5)

    # 保存
    file_name = '梁朋_Java后端_内容完善版.docx'
    try:
        doc.save(file_name)
        print(f"✅ 简历已生成：{file_name}")
        print("优化点：工作经历已补充详细职责，并与项目经历形成对应关系。")
    except Exception as e:
        print(f"❌ 生成失败：{e}")

if __name__ == "__main__":
    create_full_content_resume()

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement

class ContractFormatter:
    def __init__(self, input_path, output_path):
        self.doc = Document(input_path)
        self.output_path = output_path

        # 字号映射 (根据国家公文标准/Word标准)
        self.font_sizes = {
            '一号': 26, '二号': 22, '三号': 16, '小三': 15,
            '四号': 14, '小四': 12, '五号': 10.5
        }

        # 字体名称映射 (Windows标准名称)
        self.fonts = {
            'song': 'SimSun',          # 宋体
            'hei': 'SimHei',           # 黑体
            'fangsong': 'FangSong_GB2312', # 仿宋_GB2312 (若无则使用FangSong)
            'kai': 'KaiTi',            # 楷体
            'eng': 'Times New Roman'
        }

    def _set_font(self, run, font_name_key, size_name, bold=False):
        """通用字体设置函数"""
        font_name = self.fonts.get(font_name_key, 'SimSun')
        # 兼容处理：如果没有仿宋GB2312，尝试使用普通仿宋
        if font_name_key == 'fangsong':
            pass # 实际环境中需确认系统字体库

        run.font.name = font_name
        run.font.size = Pt(self.font_sizes[size_name])
        run.bold = bold

        # 设置中文字体 (必须使用xml设置)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def setup_page_layout(self):
        """页面设置：页边距上下2.5cm、左右3cm"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(3.0)

            # 页码设置 (python-docx原生不支持直接插入页码域，需操作XML)
            self._add_page_number(section)

    def _add_page_number(self, section):
        """在页脚添加居中的页码"""
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码域代码
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

        # 设置页码字体
        run.font.name = self.fonts['eng']
        run.font.size = Pt(self.font_sizes['小四']) # 默认跟随正文大小

    def format_cover(self):
        """
        处理封面逻辑
        注意：脚本假设文档的前几段是封面。
        需要根据实际文档结构调整索引。
        """
        # 假设第1段是合同名称
        if len(self.doc.paragraphs) > 0:
            p = self.doc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Cm(5) # 距页眉5厘米
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                self._set_font(run, 'fangsong', '一号', bold=True)

        # 假设第2段是编号 (寻找包含"编号"的段落)
        for p in self.doc.paragraphs[1:5]:
            if "编号" in p.text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # 与名称间隔两行 (约等于30-40磅)
                p.paragraph_format.space_before = Pt(30)
                for run in p.runs:
                    self._set_font(run, 'hei', '四号')
                break

        # 签约双方信息 (简单查找包含"甲方"、"乙方"的段落)
        # 实际情况可能需要更复杂的逻辑
        pass

    def format_body_structure(self):
        """核心正文格式化"""
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            # -------------------------------------------------
            # 1. 识别一级标题：合同名称 (主体部分)
            # -------------------------------------------------
            # 策略：假设封面之后，字数较少且居中的可能是正文标题
            # 这里简单演示：如果段落被标记为 Heading 1 或者用户指定逻辑
            # 由于很难自动区分封面标题和正文标题，这里仅做通用正文处理

            # -------------------------------------------------
            # 2. 识别二级标题 (条款大项，如 "第一条")
            # -------------------------------------------------
            if re.match(r'^第[一二三四五六七八九十]+条', text):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT # 虽然规范没写，但通常左对齐或缩进
                p.paragraph_format.first_line_indent = Pt(0) # 标题通常不缩进或根据需求
                for run in p.runs:
                    self._set_font(run, 'hei', '四号')
                continue

            # -------------------------------------------------
            # 3. 识别三级标题 (条款子项，如 "1." 或 "（一）")
            # -------------------------------------------------
            if re.match(r'^（[一二三四五六七八九十]+）', text) or re.match(r'^\d+\.', text):
                p.paragraph_format.first_line_indent = Pt(self.font_sizes['小四'] * 2)
                for run in p.runs:
                    self._set_font(run, 'kai', '小四')
                continue

            # -------------------------------------------------
            # 4. 基础正文 (默认为宋体小四)
            # -------------------------------------------------
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(self.font_sizes['小四'] * 2) # 首行缩进2字符
            p.paragraph_format.line_spacing = 1.25 # 行距1.25倍
            p.paragraph_format.space_after = Pt(0)

            for run in p.runs:
                # 检查是否为"以下无"或落款
                if "以下无正文" in text or "签字" in text or "盖章" in text:
                    self._set_font(run, 'song', '小四')
                    # 落款通常不需要首行缩进，可能需要整体缩进，此处保持标准
                else:
                    self._set_font(run, 'song', '小四')

    def format_tables(self):
        """表格与说明"""
        for table in self.doc.tables:
            table.autofit = False
            table.allow_autofit = False

            # 表格内容格式
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            # 判断是否为数字
                            if run.text.strip().isdigit():
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            self._set_font(run, 'song', '五号')

            # 模拟三线表 (python-docx改边框比较复杂，这里做简化处理)
            # 真正完美的三线表建议在Word模版中预设好样式，然后在这里应用样式 table.style = 'ThreeLine'
            self._set_table_borders_three_line(table)

    def _set_table_borders_three_line(self, table):
        """
        通过XML底层操作设置三线表
        顶底线粗(通常1.5pt)，中间线细(0.5pt)，无竖线
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        # 移除所有边框
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)

        # 设置顶边框 (粗)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12') # 1/8 pt units, 12 = 1.5pt
        tblBorders.append(top)

        # 设置底边框 (粗)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        tblBorders.append(bottom)

        # 设置内部横线 (细)
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:sz'), '4') # 0.5pt
        tblBorders.append(insideH)

    def run(self):
        print("正在设置页面布局...")
        self.setup_page_layout()

        print("正在格式化封面(需人工核对)...")
        self.format_cover()

        print("正在格式化正文段落...")
        self.format_body_structure()

        print("正在处理表格...")
        self.format_tables()

        print(f"保存文档至: {self.output_path}")
        self.doc.save(self.output_path)

# 使用示例
if __name__ == "__main__":
    # 请将 'input.docx' 替换为你的原始文件名
    # 确保系统中安装了 仿宋_GB2312 字体，否则可能会回退到系统默认
    formatter = ContractFormatter("raw_contract.docx", "formatted_contract.docx")
    formatter.run()
